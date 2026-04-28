from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_colored_box(doc, text, fill_color, text_color=None, bold=True):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    if text_color:
        run.font.color.rgb = RGBColor.from_string(text_color)
    cell.width = Inches(2.5)
    doc.add_paragraph()

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.footer_distance = Cm(1)
section.header_distance = Cm(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run()
run.font.name = 'Arial'
run.font.size = Pt(9)
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Benchmark Analysis Report\nAI Assistant (Башкирэнерго)')
run.font.name = 'Arial'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Date:', '2026-04-28'),
    ('Dataset:', 'combined_benchmark.csv (308 questions)'),
    ('Review method:', 'Judge LLM (deepseek-v3.2) + manual semantic verification'),
    ('', '')
]
for i, (label, value) in enumerate(meta_data):
    cell_label = meta_table.cell(i, 0)
    cell_value = meta_table.cell(i, 1)
    cell_label.paragraphs[0].add_run(label).font.name = 'Arial'
    cell_label.paragraphs[0].runs[0].font.bold = True
    cell_label.paragraphs[0].runs[0].font.size = Pt(11)
    cell_value.paragraphs[0].add_run(value).font.name = 'Arial'
    cell_value.paragraphs[0].runs[0].font.size = Pt(11)

doc.add_paragraph()

h1 = doc.add_heading('1. Executive Summary', level=1)
h1.runs[0].font.name = 'Arial'
h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

exec_summary = [
    ('Total questions:', '308'),
    ('Correct answers:', '120 (39.0%)'),
    ('Errors:', '188 (61.0%)'),
    ('Judge-model agreement:', '286/308 (92.9%)'),
    ('Answer not found:', '11 cases'),
    ('Average judge overall_score:', '3.71'),
]

metrics_table = doc.add_table(rows=len(exec_summary), cols=2)
metrics_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, value) in enumerate(exec_summary):
    cell_label = metrics_table.cell(i, 0)
    cell_value = metrics_table.cell(i, 1)
    p_label = cell_label.paragraphs[0]
    p_label.add_run(label).font.name = 'Arial'
    p_label.runs[0].font.bold = True
    p_label.runs[0].font.size = Pt(11)
    p_value = cell_value.paragraphs[0]
    p_value.add_run(value).font.name = 'Arial'
    p_value.runs[0].font.size = Pt(11)

doc.add_paragraph()

h2 = doc.add_heading('Category Breakdown', level=2)
h2.runs[0].font.name = 'Arial'

cat_table = doc.add_table(rows=5, cols=5)
cat_table.style = 'Table Grid'
cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Category', 'Questions', 'Correct', '%', 'Avg Score']
header_row = cat_table.rows[0]
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_shading(cell, 'D5E8F0')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(10)

data_rows = [
    ('Личный кабинет', '53', '29', '54.7%', '4.17'),
    ('Дополнительные услуги', '67', '22', '32.8%', '3.51'),
    ('Технологическое присоединение', '188', '69', '36.7%', '3.65'),
    ('Всего', '308', '120', '39.0%', '3.71'),
]
for row_idx, row_data in enumerate(data_rows, 1):
    for col_idx, text in enumerate(row_data):
        cell = cat_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if row_idx == 4:
            run.font.bold = True

doc.add_paragraph()

h1 = doc.add_heading('2. Detailed Statistics by Category', level=1)
h1.runs[0].font.name = 'Arial'
h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

h2 = doc.add_heading('2.1. Личный кабинет (ЛК)', level=2)
h2.runs[0].font.name = 'Arial'

p = doc.add_paragraph()
p.add_run('Total: 53 questions | Correct: 29 (54.7%) | Errors: 24 (45.3%)').font.name = 'Arial'

lk_table = doc.add_table(rows=5, cols=3)
lk_table.style = 'Table Grid'
lk_table.alignment = WD_TABLE_ALIGNMENT.LEFT
lk_headers = ['Classification', 'Count', '%']
for i, header in enumerate(lk_headers):
    cell = lk_table.rows[0].cells[i]
    set_cell_shading(cell, 'D5E8F0')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(10)

lk_data = [
    ('Exact match', '17', '32.1%'),
    ('Meaning correct', '12', '22.6%'),
    ('Error', '21', '39.6%'),
    ('Not found', '3', '5.7%'),
]
for row_idx, row_data in enumerate(lk_data, 1):
    for col_idx, text in enumerate(row_data):
        cell = lk_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

doc.add_paragraph()

h2 = doc.add_heading('2.2. Дополнительные услуги (ДУ)', level=2)
h2.runs[0].font.name = 'Arial'

p = doc.add_paragraph()
p.add_run('Total: 67 questions | Correct: 22 (32.8%) | Errors: 45 (67.2%)').font.name = 'Arial'

du_table = doc.add_table(rows=4, cols=3)
du_table.style = 'Table Grid'
du_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, header in enumerate(lk_headers):
    cell = du_table.rows[0].cells[i]
    set_cell_shading(cell, 'D5E8F0')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(10)

du_data = [
    ('Exact match', '9', '13.4%'),
    ('Meaning correct', '13', '19.4%'),
    ('Error', '45', '67.2%'),
]
for row_idx, row_data in enumerate(du_data, 1):
    for col_idx, text in enumerate(row_data):
        cell = du_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

doc.add_paragraph()

h2 = doc.add_heading('2.3. Технологическое присоединение (ТПП)', level=2)
h2.runs[0].font.name = 'Arial'

p = doc.add_paragraph()
p.add_run('Total: 188 questions | Correct: 69 (36.7%) | Errors: 119 (63.3%)').font.name = 'Arial'

tpp_table = doc.add_table(rows=5, cols=3)
tpp_table.style = 'Table Grid'
tpp_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, header in enumerate(lk_headers):
    cell = tpp_table.rows[0].cells[i]
    set_cell_shading(cell, 'D5E8F0')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(10)

tpp_data = [
    ('Exact match', '38', '20.2%'),
    ('Meaning correct', '31', '16.5%'),
    ('Error', '111', '59.0%'),
    ('Not found', '8', '4.3%'),
]
for row_idx, row_data in enumerate(tpp_data, 1):
    for col_idx, text in enumerate(row_data):
        cell = tpp_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

doc.add_paragraph()

h1 = doc.add_heading('3. Judge LLM Objectivity Analysis', level=1)
h1.runs[0].font.name = 'Arial'
h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
run = p.add_run('The judge model (deepseek-v3.2) achieved ')
run.font.name = 'Arial'
run = p.add_run('286/308 (92.9%)')
run.font.name = 'Arial'
run.font.bold = True
run = p.add_run(' agreement with manual semantic review.')
run.font.name = 'Arial'

doc.add_paragraph()

h2 = doc.add_heading('Disagreements (5 cases):', level=2)
h2.runs[0].font.name = 'Arial'

disagreements = [
    '#43 ДУ: Judge=0, Manual=1, Judge says wrong but justification positive',
    '#91 ДУ: Judge=0, Manual=1, Judge says wrong but justification positive',
    '#94 ДУ: Judge=0, Manual=1, Judge says wrong but justification positive',
    '#99 ДУ: Judge=0, Manual=1, Judge says wrong but justification positive',
    '#111 ТПП: Judge=0, Manual=1, Judge says wrong but justification positive',
]
for item in disagreements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

doc.add_paragraph()

h2 = doc.add_heading('Judge bias analysis:', level=2)
h2.runs[0].font.name = 'Arial'

bias_data = [
    ('Judge was too strict:', '5'),
    ('Judge was too lenient:', '0'),
    ('No judge data available:', '17'),
]
for label, value in bias_data:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.font.name = 'Arial'
    run.font.bold = True
    run = p.add_run(value)
    run.font.name = 'Arial'

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Bias verdict: ')
run.font.name = 'Arial'
run.font.bold = True
run = p.add_run('Judge tends to be overly strict')
run.font.name = 'Arial'
run.font.italic = True

doc.add_paragraph()

h1 = doc.add_heading('4. Error Patterns (Anna Borisovna)', level=1)
h1.runs[0].font.name = 'Arial'
h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

error_table = doc.add_table(rows=7, cols=3)
error_table.style = 'Table Grid'
error_table.alignment = WD_TABLE_ALIGNMENT.LEFT
error_headers = ['Error Pattern', 'Occurrences', '% of Errors']
for i, header in enumerate(error_headers):
    cell = error_table.rows[0].cells[i]
    set_cell_shading(cell, 'D5E8F0')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.size = Pt(10)

error_data = [
    ('Incorrect terminology', '69', '36.7%'),
    ('Power/category limitations not considered', '41', '21.8%'),
    ('Incorrect cost calculation or tariffs', '43', '22.9%'),
    ('Social benefit rate applied incorrectly', '22', '11.7%'),
    ('Confusion between ДУ and ТП', '6', '3.2%'),
    ('TPP procedure order violation', '0', '0.0%'),
]
for row_idx, row_data in enumerate(error_data, 1):
    for col_idx, text in enumerate(row_data):
        cell = error_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Total errors matching patterns: ')
run.font.name = 'Arial'
run = p.add_run('181')
run.font.name = 'Arial'
run.font.bold = True
run = p.add_run(' across ')
run.font.name = 'Arial'
run = p.add_run('115')
run.font.name = 'Arial'
run.font.bold = True
run = p.add_run(' questions')
run.font.name = 'Arial'

doc.add_paragraph()

h1 = doc.add_heading('5. Development Roadmap', level=1)
h1.runs[0].font.name = 'Arial'
h1.runs[0].font.color.rgb = RGBColor(0, 51, 102)

h2 = doc.add_heading('Immediate Actions (Week 1-2)', level=2)
h2.runs[0].font.name = 'Arial'

immediate = [
    'Fix ДУ category (32.8% accuracy): Model confuses ДУ procedures with ТП. Need separate prompt for ДУ questions.',
    'Improve search quality for deadlines/costs: Model frequently gives wrong numbers. Add structured data tables to vector DB.',
    'Fix password recovery procedures: Model applies generic recovery to all cases, ignoring key context.',
]
for item in immediate:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

doc.add_paragraph()

h2 = doc.add_heading('Short-term (Week 3-4)', level=2)
h2.runs[0].font.name = 'Arial'

shortterm = [
    'Add Anna corrections to prompts: Incorporate domain expert feedback into system prompt.',
    'Improve source attribution: Model says \'not found\' 11 times where answer exists. Fix search recall.',
    'Power/category constraint awareness: Add explicit rules about ФЛ≤15kW, ИП/ЮЛ≤150kW.',
]
for item in shortterm:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

doc.add_paragraph()

h2 = doc.add_heading('Medium-term (Month 2-3)', level=2)
h2.runs[0].font.name = 'Arial'

mediumterm = [
    'Separate response agents per category: Different system prompts for ЛК/ДУ/ТПП.',
    'RAG quality monitoring: Track context_recall scores, alert on degradation.',
    'Benchmark automation: Run weekly benchmark runs, track accuracy trends.',
    'Expand benchmark: Add more ДУ-specific and edge case questions.',
]
for item in mediumterm:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

doc.save('D:/ai_assistant/benchmark_report_generated.docx')
print('Report saved to: D:/ai_assistant/benchmark_report_generated.docx')
